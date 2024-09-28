import streamlit as st
from multiprocessing import Pool, cpu_count
import fitz
import tempfile
import subprocess
from pathlib import Path
import pytesseract
from PIL import Image
from io import BytesIO
from docx import Document
import docx
import csv
import os

# Initialize environment variables for Tesseract if needed
pytesseract.pytesseract.tesseract_cmd = os.getenv("LAMBDA_TASK_ROOT", "") + "/bin/tesseract"
os.environ['TESSDATA_PREFIX'] = os.getenv("LAMBDA_TASK_ROOT", "") + "/tesseract/share/tessdata"
os.environ['LD_LIBRARY_PATH'] = os.getenv("LAMBDA_TASK_ROOT", "") + "/lib"




def process_docx(file):
    doc = docx.Document(file)

    # Extract text from paragraphs and tables
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            full_text.append("\t".join(row_data))  # Join cell data with a tab for better readability


    return "\n".join(full_text)
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            full_text.append("\t".join(row_data))  # Join cell data with a tab for better readability

    return "\n".join(full_text)



def convert_doc_to_docx(doc_file_path, output_dir):
    if not os.path.isfile(doc_file_path):
        raise FileNotFoundError(f"The file {doc_file_path} does not exist.")

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    base_name = os.path.basename(doc_file_path)
    docx_file_name = os.path.splitext(base_name)[0] + '.docx'
    output_file_path = os.path.join(output_dir, docx_file_name)

    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--outdir', output_dir, doc_file_path], check=True)
        print(f"Conversion successful: {doc_file_path} to {output_file_path}")
        return output_file_path
    except FileNotFoundError:
        print("LibreOffice is not installed or not found in PATH.")
    except subprocess.CalledProcessError as e:
        print(f"An error occurred during conversion: {e}")



def extract_text_from_pages_single_threaded(pdf_path):
    extracted_text = ""
    doc = fitz.open(pdf_path)
    num_pages = doc.page_count
    for i in range(num_pages):
        page = doc.load_page(i)
        extracted_text += page.get_text("text")
        if not extracted_text.strip():
            extracted_text = extract_text_with_tesseract(pdf_path)
        extracted_text += f"\n--- End of Page {i + 1} ---\n"
    return extracted_text

def extract_text_from_page_indices(pdf_path, indices):
    extracted_text = ""
    doc = fitz.open(pdf_path)
    for i in indices:
        page = doc.load_page(i)
        extracted_text += page.get_text("text")
        if not extracted_text.strip():
            extracted_text = extract_text_with_tesseract(pdf_path, pages=[i])
        extracted_text += f"\n--- End of Page {i + 1} ---\n"
    return extracted_text

def extract_text_with_tesseract(pdf_path, pages=None):
    extracted_text = ""
    doc = fitz.open(pdf_path)
    page_range = range(doc.page_count) if pages is None else pages
    for i in page_range:
        page = doc.load_page(i)
        pix = page.get_pixmap()
        image = Image.open(BytesIO(pix.tobytes(output="png")))
        extracted_text += pytesseract.image_to_string(image)
        extracted_text += f"\n--- End of Page {i + 1} ---\n"
    return extracted_text

def parallel_text_extraction(pdf_path, num_pages):
    cpu = cpu_count()
    seg_size = int(num_pages / cpu + 1)
    indices = [range(i * seg_size, min((i + 1) * seg_size, num_pages)) for i in range(cpu)]
    with Pool() as pool:
        results = pool.starmap(extract_text_from_page_indices, [(pdf_path, idx) for idx in indices])
    combined_text = "".join(results)
    return combined_text

def process_pdf(file_path):
    doc = fitz.open(file_path)
    num_pages = doc.page_count
    if num_pages < 10:
        extracted_text = extract_text_from_pages_single_threaded(file_path)
    else:
        extracted_text = parallel_text_extraction(file_path, num_pages)
    return extracted_text

def process_csv(csv_path):
    with open(csv_path, 'r') as f:
        reader = csv.reader(f)
        return "\n".join([",".join(row) for row in reader])

# Streamlit App
st.title("Document Processing App")

# File upload
uploaded_file = st.file_uploader("Upload a PDF, DOCX, DOC, or CSV file", type=["pdf", "docx", "doc", "csv"])

if uploaded_file is not None:
    # Process based on file type
    file_type = uploaded_file.name.split('.')[-1]
    
    if file_type == "pdf":
        with open(uploaded_file.name, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.write("Extracting text from PDF...")
        extracted_text = process_pdf(uploaded_file.name)
        st.text_area("Extracted Text", value=extracted_text, height=300)
    
    elif file_type == "docx":
        with open(uploaded_file.name, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.write("Extracting text from DOCX...")
        extracted_text = process_docx(uploaded_file.name)
        st.text_area("Extracted Text", value=extracted_text, height=300)

    elif file_type == "doc":
        with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_doc_file:
            temp_doc_file.write(uploaded_file.getbuffer())
            temp_doc_file.flush()

            output_dir = tempfile.mkdtemp()
            docx_file_path = convert_doc_to_docx(temp_doc_file.name, output_dir)

            if docx_file_path:
                st.write("Extracting text from DOC...")
                extracted_text = process_docx(docx_file_path)
                st.text_area("Extracted Text", value=extracted_text, height=300)

            # Clean up temporary files
            os.remove(temp_doc_file.name)
            os.remove(docx_file_path)

    elif file_type == "csv":
        st.write("Processing CSV...")
        extracted_text = process_csv(uploaded_file)
        st.text_area("CSV Content", value=extracted_text, height=300)
